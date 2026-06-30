"""ATS-quality CV generator — DOCX, PDF (single-column), and JSON-Resume export.

Takes a profile dict (as returned by ``_profile_to_payload()`` or equivalent)
and produces downloadable artefacts. All three formats share the same section
ordering and are designed to round-trip cleanly through the CV parser.
"""
from __future__ import annotations

import io
import json
from typing import Any

from app.core.logging import get_logger

logger = get_logger(__name__)


# ── helpers ───────────────────────────────────────────────────────────────────

def _s(value: Any, default: str = "") -> str:
    return str(value or default).strip()


def _list(value: Any) -> list[str]:
    if isinstance(value, list):
        return [str(v).strip() for v in value if str(v).strip()]
    if isinstance(value, str) and value.strip():
        return [v.strip() for v in value.split(",") if v.strip()]
    return []


def _fmt_range(start: str, end: str, current: bool = False) -> str:
    import re as _re

    def _to_display(d: str) -> str:
        d = _s(d)
        # YYYY-MM → MM/YYYY  (European format; also parseable by the CV parser's _DATE_RANGE regex)
        if _re.match(r"^\d{4}-\d{2}$", d):
            return f"{d[5:]}/{d[:4]}"
        return d

    s = _to_display(start)
    if current or _s(end).lower() in ("present", "presente", "atual", ""):
        e = "Presente"
    else:
        e = _to_display(end)
    if not s and not e:
        return ""
    if not s:
        return e
    if not e or s == e:
        return s
    return f"{s} – {e}"


# ── JSON-Resume export ────────────────────────────────────────────────────────

def to_json_resume(profile: dict[str, Any]) -> dict[str, Any]:
    """Map profile dict to the JSON-Resume v1 schema (jsonresume.org)."""
    full_name = _s(profile.get("fullName") or profile.get("full_name"))
    name_parts = full_name.split()
    first = name_parts[0] if name_parts else ""
    last = " ".join(name_parts[1:]) if len(name_parts) > 1 else ""

    basics: dict[str, Any] = {
        "name": full_name,
        "label": _s(profile.get("professionalTitle") or profile.get("jobTitle") or profile.get("job_title")),
        "email": _s(profile.get("email")),
        "phone": _s(profile.get("phone")),
        "summary": _s(profile.get("professionalSummary") or profile.get("professional_summary") or profile.get("summary")),
        "location": {"address": _s(profile.get("location"))},
        "profiles": [],
    }
    if linkedin := _s(profile.get("linkedinUrl") or profile.get("linkedin_url")):
        basics["profiles"].append({"network": "LinkedIn", "url": linkedin})
    if github := _s(profile.get("githubUrl") or profile.get("github_url")):
        basics["profiles"].append({"network": "GitHub", "url": github})
    if portfolio := _s(profile.get("portfolioUrl") or profile.get("portfolio_url")):
        basics["profiles"].append({"network": "Portfolio", "url": portfolio})

    work = []
    for exp in _list_of_dicts(profile.get("workExperience") or profile.get("work_experience") or profile.get("experience")):
        start = _s(exp.get("startDate") or exp.get("start_date"))
        end = _s(exp.get("endDate") or exp.get("end_date"))
        work.append({
            "name": _s(exp.get("company")),
            "position": _s(exp.get("jobTitle") or exp.get("role")),
            "location": _s(exp.get("location")),
            "startDate": start,
            "endDate": "" if (exp.get("current") or end.lower() in ("present", "presente", "atual")) else end,
            "summary": _s(exp.get("description")),
        })

    edu = []
    for e in _list_of_dicts(profile.get("education")):
        start = _s(e.get("startDate") or e.get("start_date"))
        end = _s(e.get("endDate") or e.get("end_date"))
        edu.append({
            "institution": _s(e.get("institution")),
            "area": _s(e.get("fieldOfStudy") or e.get("field_of_study")),
            "studyType": _s(e.get("degree")),
            "startDate": start,
            "endDate": end,
        })

    all_skills = (
        _list(profile.get("hardSkills") or profile.get("hard_skills"))
        + _list(profile.get("techniques"))
        + _list(profile.get("tools"))
    )
    if not all_skills:
        all_skills = _list(profile.get("skills"))

    skills_section = []
    if _list(profile.get("hardSkills") or profile.get("hard_skills")):
        skills_section.append({"name": "Hard Skills", "keywords": _list(profile.get("hardSkills") or profile.get("hard_skills"))})
    if _list(profile.get("techniques")):
        skills_section.append({"name": "Techniques", "keywords": _list(profile.get("techniques"))})
    if _list(profile.get("tools")):
        skills_section.append({"name": "Tools", "keywords": _list(profile.get("tools"))})
    if not skills_section and all_skills:
        skills_section.append({"name": "Skills", "keywords": all_skills})

    languages_section = [
        {"language": lng, "fluency": ""}
        for lng in _list(profile.get("languages"))
    ]

    return {
        "$schema": "https://raw.githubusercontent.com/jsonresume/resume-schema/v1.0.0/schema.json",
        "basics": basics,
        "work": work,
        "education": edu,
        "skills": skills_section,
        "languages": languages_section,
        "certificates": [{"name": c} for c in _list(profile.get("certifications"))],
        "meta": {"theme": "parvagas"},
    }


# ── DOCX export ───────────────────────────────────────────────────────────────

def to_docx(profile: dict[str, Any]) -> bytes:
    """Generate a single-column ATS-friendly DOCX from profile dict."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError as exc:
        raise RuntimeError("python-docx is required for DOCX export") from exc

    doc = Document()

    # Page margins: narrow for more content space.
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # Remove default paragraph spacing.
    from docx.oxml.ns import qn as _qn
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    def _add_hr(para):
        """Add a bottom border (HR) to a paragraph via XML."""
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "CCCCCC")
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _section_heading(text: str):
        p = doc.add_paragraph()
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)  # dark red
        _add_hr(p)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        return p

    def _entry_header(left: str, right: str = ""):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(1)
        if right:
            # left-aligned name, right-aligned date via tab stop
            tab = OxmlElement("w:tab")
            p._p.get_or_add_pPr()
            run_l = p.add_run(left)
            run_l.bold = True
            run_l.font.size = Pt(10)
            p.add_run("\t")
            run_r = p.add_run(right)
            run_r.font.size = Pt(9)
            run_r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            # Set a right-side tab stop at ~14 cm (safe for A4 with 2cm margins).
            from docx.oxml import OxmlElement as OE
            pPr = p._p.get_or_add_pPr()
            tabs = OE("w:tabs")
            tab_el = OE("w:tab")
            tab_el.set(qn("w:val"), "right")
            tab_el.set(qn("w:pos"), "8640")  # twips: ~15.2 cm
            tabs.append(tab_el)
            pPr.append(tabs)
        else:
            run = p.add_run(left)
            run.bold = True
            run.font.size = Pt(10)
        return p

    def _small(text: str, italic: bool = False, color: tuple = (0x44, 0x44, 0x44)):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(text)
        run.font.size = Pt(9)
        run.italic = italic
        run.font.color.rgb = RGBColor(*color)
        return p

    def _bullet(text: str):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Inches(0.2)
        run = p.add_run(text)
        run.font.size = Pt(9)
        return p

    # ── Header ──────────────────────────────────────────────────────────────
    full_name = _s(profile.get("fullName") or profile.get("full_name"))
    title = _s(profile.get("professionalTitle") or profile.get("jobTitle") or profile.get("job_title"))
    email = _s(profile.get("email"))
    phone = _s(profile.get("phone"))
    location = _s(profile.get("location"))
    linkedin = _s(profile.get("linkedinUrl") or profile.get("linkedin_url"))

    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(full_name or "Nome do Candidato")
    name_run.bold = True
    name_run.font.size = Pt(18)
    name_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    if title:
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(title)
        title_run.font.size = Pt(11)
        title_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    contact_parts = [p for p in [location, phone, email, linkedin] if p]
    if contact_parts:
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_para.paragraph_format.space_after = Pt(6)
        contact_run = contact_para.add_run("  |  ".join(contact_parts))
        contact_run.font.size = Pt(9)
        contact_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    summary = _s(profile.get("professionalSummary") or profile.get("professional_summary") or profile.get("summary"))
    if summary:
        _section_heading("Resumo Profissional")
        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(2)
        sp.paragraph_format.space_after = Pt(4)
        run = sp.add_run(summary)
        run.font.size = Pt(9)

    # ── Experience ───────────────────────────────────────────────────────────
    experience = _list_of_dicts(profile.get("workExperience") or profile.get("work_experience") or profile.get("experience"))
    if experience:
        _section_heading("Experiência Profissional")
        for exp in experience:
            company = _s(exp.get("company"))
            job_title = _s(exp.get("jobTitle") or exp.get("role"))
            loc = _s(exp.get("location"))
            date_str = _fmt_range(
                _s(exp.get("startDate") or exp.get("start_date")),
                _s(exp.get("endDate") or exp.get("end_date")),
                bool(exp.get("current")),
            )
            company_loc = f"{company}, {loc}" if loc and company else (company or loc)
            if company_loc:
                _entry_header(company_loc, date_str)
            if job_title:
                _small(job_title, italic=True)
            desc = _s(exp.get("description"))
            if desc:
                for bullet_text in desc.split(". "):
                    if bullet_text.strip():
                        _bullet(bullet_text.strip().rstrip(".") + ".")

    # ── Education ────────────────────────────────────────────────────────────
    education = _list_of_dicts(profile.get("education"))
    if education:
        _section_heading("Formação Académica")
        for edu in education:
            institution = _s(edu.get("institution"))
            degree = _s(edu.get("degree"))
            field = _s(edu.get("fieldOfStudy") or edu.get("field_of_study"))
            loc = _s(edu.get("location"))
            date_str = _fmt_range(
                _s(edu.get("startDate") or edu.get("start_date")),
                _s(edu.get("endDate") or edu.get("end_date")),
            )
            inst_loc = f"{institution}, {loc}" if loc and institution else (institution or loc)
            if inst_loc:
                _entry_header(inst_loc, date_str)
            degree_field = " – ".join(p for p in [degree, field] if p)
            if degree_field:
                _small(degree_field, italic=True)

    # ── Skills ───────────────────────────────────────────────────────────────
    hard = _list(profile.get("hardSkills") or profile.get("hard_skills"))
    techniques = _list(profile.get("techniques"))
    tools = _list(profile.get("tools"))
    flat_skills = _list(profile.get("skills"))
    # Prefer split buckets; fall back to flat list.
    if hard or techniques or tools:
        _section_heading("Competências")
        if hard:
            _small(f"Hard Skills: {', '.join(hard)}")
        if techniques:
            _small(f"Técnicas: {', '.join(techniques)}")
        if tools:
            _small(f"Ferramentas: {', '.join(tools)}")
    elif flat_skills:
        _section_heading("Competências")
        _small(", ".join(flat_skills))

    # ── Languages ────────────────────────────────────────────────────────────
    languages = _list(profile.get("languages"))
    if languages:
        _section_heading("Idiomas")
        _small(", ".join(languages))

    # ── Certifications ────────────────────────────────────────────────────────
    certs = _list(profile.get("certifications"))
    if certs:
        _section_heading("Certificações")
        for cert in certs:
            _small(f"• {cert}")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── PDF export ────────────────────────────────────────────────────────────────

def to_pdf(profile: dict[str, Any]) -> bytes:
    """Generate a single-column ATS-friendly PDF from profile dict."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import mm
        from reportlab.lib.colors import HexColor, Color
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
    except ImportError as exc:
        raise RuntimeError("reportlab is required for PDF export") from exc

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=20 * mm,
        rightMargin=20 * mm,
        topMargin=15 * mm,
        bottomMargin=15 * mm,
    )

    styles = getSampleStyleSheet()
    dark = HexColor("#1a1a2e")
    red = HexColor("#8B0000")
    gray = HexColor("#555555")
    light_gray = HexColor("#888888")

    # Explicit leading + generous spaceAfter on the name: a 20pt Helvetica-Bold
    # glyph box is ~27.6pt tall, which exceeds reportlab's default leading (24pt).
    # Without the extra spacing the title overlaps the name's descenders.
    st_name = ParagraphStyle("Name", parent=styles["Normal"], fontSize=20, leading=24, textColor=dark,
                             fontName="Helvetica-Bold", alignment=TA_CENTER, spaceAfter=10)
    st_title = ParagraphStyle("Title", parent=styles["Normal"], fontSize=11, leading=14, textColor=gray,
                              alignment=TA_CENTER, spaceAfter=2)
    st_contact = ParagraphStyle("Contact", parent=styles["Normal"], fontSize=8.5, textColor=gray,
                                alignment=TA_CENTER, spaceAfter=8)
    st_section = ParagraphStyle("Section", parent=styles["Normal"], fontSize=10, textColor=red,
                                fontName="Helvetica-Bold", spaceBefore=10, spaceAfter=3)
    st_entry = ParagraphStyle("Entry", parent=styles["Normal"], fontSize=10,
                              fontName="Helvetica-Bold", spaceBefore=4, spaceAfter=1)
    st_sub = ParagraphStyle("Sub", parent=styles["Normal"], fontSize=9, textColor=gray,
                            fontName="Helvetica-Oblique", spaceAfter=1)
    st_body = ParagraphStyle("Body", parent=styles["Normal"], fontSize=9, spaceAfter=1,
                             leading=13)
    st_bullet = ParagraphStyle("Bullet", parent=styles["Normal"], fontSize=9,
                               leftIndent=12, spaceAfter=1, leading=13,
                               firstLineIndent=0, bulletIndent=4)

    elements = []

    def hr():
        elements.append(HRFlowable(width="100%", thickness=0.5, color=HexColor("#CCCCCC"), spaceAfter=4))

    full_name = _s(profile.get("fullName") or profile.get("full_name"))
    title = _s(profile.get("professionalTitle") or profile.get("jobTitle") or profile.get("job_title"))
    email = _s(profile.get("email"))
    phone = _s(profile.get("phone"))
    location = _s(profile.get("location"))
    linkedin = _s(profile.get("linkedinUrl") or profile.get("linkedin_url"))

    elements.append(Paragraph(full_name or "Nome do Candidato", st_name))
    if title:
        elements.append(Paragraph(title, st_title))
    contact_parts = [p for p in [location, phone, email, linkedin] if p]
    if contact_parts:
        elements.append(Paragraph("  |  ".join(contact_parts), st_contact))

    summary = _s(profile.get("professionalSummary") or profile.get("professional_summary") or profile.get("summary"))
    if summary:
        elements.append(Paragraph("RESUMO PROFISSIONAL", st_section))
        hr()
        elements.append(Paragraph(summary, st_body))

    experience = _list_of_dicts(profile.get("workExperience") or profile.get("work_experience") or profile.get("experience"))
    if experience:
        elements.append(Paragraph("EXPERIÊNCIA PROFISSIONAL", st_section))
        hr()
        for exp in experience:
            company = _s(exp.get("company"))
            job_title = _s(exp.get("jobTitle") or exp.get("role"))
            loc = _s(exp.get("location"))
            date_str = _fmt_range(
                _s(exp.get("startDate") or exp.get("start_date")),
                _s(exp.get("endDate") or exp.get("end_date")),
                bool(exp.get("current")),
            )
            company_loc = f"{company}, {loc}" if loc and company else (company or loc)
            header = f"{company_loc}  –  {date_str}" if date_str else company_loc
            if header:
                elements.append(Paragraph(header, st_entry))
            if job_title:
                elements.append(Paragraph(job_title, st_sub))
            desc = _s(exp.get("description"))
            if desc:
                for bullet_text in desc.split(". "):
                    if bullet_text.strip():
                        elements.append(Paragraph(f"• {bullet_text.strip().rstrip('.')}.", st_bullet))

    education = _list_of_dicts(profile.get("education"))
    if education:
        elements.append(Paragraph("FORMAÇÃO ACADÉMICA", st_section))
        hr()
        for edu in education:
            institution = _s(edu.get("institution"))
            degree = _s(edu.get("degree"))
            field = _s(edu.get("fieldOfStudy") or edu.get("field_of_study"))
            loc = _s(edu.get("location"))
            date_str = _fmt_range(
                _s(edu.get("startDate") or edu.get("start_date")),
                _s(edu.get("endDate") or edu.get("end_date")),
            )
            inst_loc = f"{institution}, {loc}" if loc and institution else (institution or loc)
            header = f"{inst_loc}  –  {date_str}" if date_str else inst_loc
            if header:
                elements.append(Paragraph(header, st_entry))
            degree_field = " – ".join(p for p in [degree, field] if p)
            if degree_field:
                elements.append(Paragraph(degree_field, st_sub))

    hard = _list(profile.get("hardSkills") or profile.get("hard_skills"))
    techniques = _list(profile.get("techniques"))
    tools = _list(profile.get("tools"))
    flat_skills = _list(profile.get("skills"))
    if hard or techniques or tools:
        elements.append(Paragraph("COMPETÊNCIAS", st_section))
        hr()
        if hard:
            elements.append(Paragraph(f"<b>Hard Skills:</b> {', '.join(hard)}", st_body))
        if techniques:
            elements.append(Paragraph(f"<b>Técnicas:</b> {', '.join(techniques)}", st_body))
        if tools:
            elements.append(Paragraph(f"<b>Ferramentas:</b> {', '.join(tools)}", st_body))
    elif flat_skills:
        elements.append(Paragraph("COMPETÊNCIAS", st_section))
        hr()
        elements.append(Paragraph(", ".join(flat_skills), st_body))

    languages = _list(profile.get("languages"))
    if languages:
        elements.append(Paragraph("IDIOMAS", st_section))
        hr()
        elements.append(Paragraph(", ".join(languages), st_body))

    certs = _list(profile.get("certifications"))
    if certs:
        elements.append(Paragraph("CERTIFICAÇÕES", st_section))
        hr()
        for cert in certs:
            elements.append(Paragraph(f"• {cert}", st_bullet))

    doc.build(elements)
    return buf.getvalue()


# ── helpers ───────────────────────────────────────────────────────────────────

def _list_of_dicts(value: Any) -> list[dict]:
    if isinstance(value, list):
        return [v for v in value if isinstance(v, dict)]
    if isinstance(value, str):
        try:
            parsed = json.loads(value)
            if isinstance(parsed, list):
                return [v for v in parsed if isinstance(v, dict)]
        except Exception:
            pass
    return []
